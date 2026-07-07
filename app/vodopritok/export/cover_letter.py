"""Сопроводительное письмо и чек-лист ТЗ для отклика."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from vodopritok.demo.session import TZ_MAPPING
from vodopritok.models import OUTPUT_DIR, ReservoirCard


def generate_cover_letter(
    reservoir: ReservoirCard,
    expert: str = "Эксперт по ОВП",
    company: str = "Заказчик",
    path: Path | None = None,
) -> Path:
    out = path or OUTPUT_DIR / "00-soprovoditelnoe-pismo.docx"
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
    doc.add_paragraph()
    doc.add_paragraph("Уважаемые коллеги!")
    doc.add_paragraph()
    doc.add_paragraph(
        f"Направляю материалы по проекту селективного ограничения водопритока для "
        f"{reservoir.field_name}. Предлагаю ведение R&D как эксперт (Physics+ML, хемоинформатика) "
        f"с воспроизводимым in silico конвейером 500→top-5 и полным пакетом deliverables по ТЗ за 4 месяца."
    )
    doc.add_paragraph()
    doc.add_paragraph("В приложении:")
    for item in [
        "One-pager с рекомендацией трека и top-5 под ваш пласт",
        "6 документов .docx по пунктам ТЗ",
        "Демо-данные screening и QSPRpred validation",
        "План снижения OPEX: 10 мероприятий, методология, NPV/payback (/opex-plan)",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "Экономическая компетенция: понимание CAPEX/OPEX внедрения новых технологий ОВП, "
        "разработка мероприятий по сокращению операционных затрат (закупки, ремонт, мониторинг WC), "
        "обоснование payback/NPV до выхода в ОПР."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Метод: механизм обводнения → трек технологии → in silico screening → lab gate → ОПР. "
        "Lab gate (Frrw≥5, Frro≤2) — обязательный фильтр перед выходом в поле."
    )
    doc.add_paragraph()
    doc.add_paragraph("Готов обсудить live demo и адаптацию под параметры вашего пласта.")
    doc.add_paragraph()
    doc.add_paragraph(expert)

    doc.save(out)
    return out


def generate_tz_checklist(path: Path | None = None) -> Path:
    out = path or OUTPUT_DIR / "00-checklist-tz.docx"
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("Чек-лист покрытия ТЗ проекта", level=1)
    table = doc.add_table(rows=1 + len(TZ_MAPPING), cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Пункт ТЗ", "Deliverable", "Статус", "Как показываю"]):
        hdr[i].text = h
    for ri, row in enumerate(TZ_MAPPING):
        cells = table.rows[ri + 1].cells
        cells[0].text = row["tz"]
        cells[1].text = row["deliverable"]
        cells[2].text = "✓ Demo / docx"
        cells[3].text = row["demo"]

    doc.save(out)
    return out
