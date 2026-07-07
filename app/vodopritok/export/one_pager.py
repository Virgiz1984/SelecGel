"""One-pager для заказчика после screening."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from vodopritok.decision_tree import TechnologyRecommendation, get_primary_tracks
from vodopritok.economics import OpexScenario, analyze_opex, build_opex_plan
from vodopritok.models import OUTPUT_DIR, ProjectContext, ReservoirCard


def generate_one_pager(
    reservoir: ReservoirCard,
    recommendations: list[TechnologyRecommendation],
    top5: list[dict],
    expert: str = "Эксперт по ОВП",
    company: str = "Заказчик",
    path: Path | None = None,
) -> Path:
    out = path or OUTPUT_DIR / "00-one-pager-rekomendaciya.docx"
    out.parent.mkdir(parents=True, exist_ok=True)

    wc_after = max(reservoir.water_cut_pct - 18, 55)
    opex = analyze_opex(OpexScenario(
        name=reservoir.field_name,
        water_cut_before_pct=reservoir.water_cut_pct,
        water_cut_after_pct=wc_after,
        oil_rate_tpd=reservoir.oil_rate_tpd,
    ))

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Рекомендация по селективному ОВП — one-pager")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(
        f"{company} · {reservoir.field_name} · {reservoir.well_name or 'скважина'} · "
        f"{datetime.now().strftime('%d.%m.%Y')}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Эксперт: {expert}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading("Карточка пласта", level=2)
    doc.add_paragraph(
        f"T={reservoir.temperature_c}°C · salinity {reservoir.salinity_g_l} g/L · "
        f"k={reservoir.permeability_md} mD · WC={reservoir.water_cut_pct}% · "
        f"Qn={reservoir.oil_rate_tpd} t/d · механизм: {reservoir.water_mechanism}"
    )

    doc.add_heading("Стратегия R&D (двухтрековая)", level=2)
    tracks = get_primary_tracks(reservoir)
    if tracks["track1"]:
        doc.add_paragraph(
            f"Track 1 (primary): {tracks['track1'].name_ru} — in silico top-5 ниже, "
            "синтез 5–8 полимеров, core flood Frrw/Frro."
        )
    if tracks["track2"]:
        doc.add_paragraph(
            f"Track 2 (backup): {tracks['track2'].name_ru} — 2–3 композиции параллельно; "
            "активируется если gate RPM не достигнут или нужен HTHS-запас."
        )
    if tracks["track3"]:
        doc.add_paragraph(
            f"Track 3 (условный): {tracks['track3'].name_ru} — при канальном обводнении."
        )

    doc.add_heading("Рекомендуемые технологии (ranking)", level=2)
    for rec in recommendations[:3]:
        p = doc.add_paragraph()
        p.add_run(f"{rec.rank}. {rec.name_ru} ({rec.track}, score={rec.score:.0f})").bold = True
        doc.add_paragraph(rec.rationale[:200])

    doc.add_heading("Top-5 in silico (lab gate)", level=2)
    table = doc.add_table(rows=1 + len(top5), cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["#", "mol_id", "Frrw", "Frro"]):
        hdr[i].text = h
    for ri, m in enumerate(top5):
        cells = table.rows[ri + 1].cells
        cells[0].text = str(m.get("rank", ri + 1))
        cells[1].text = str(m.get("mol_id", ""))
        cells[2].text = f"{float(m.get('predicted_frrw', 0)):.2f}"
        cells[3].text = f"{float(m.get('predicted_frro', 0)):.2f}"

    doc.add_heading("Экономика (прогноз)", level=2)
    doc.add_paragraph(
        f"WC после ОВP: {wc_after:.0f}% · экономия воды: {opex.water_reduction_m3_year:,.0f} m³/год · "
        f"годовая экономия: {opex.annual_savings_rub:,.0f} ₽ · payback: {opex.payback_months:.0f} мес · "
        f"NPV 3 года: {opex.npv_3yr_rub:,.0f} ₽"
    )

    plan = build_opex_plan(ProjectContext(reservoir=reservoir))
    doc.add_heading("Мероприятия по сокращению OPEX (топ-5)", level=2)
    top_measures = sorted(plan["measures"], key=lambda m: m["estimated_annual_rub"], reverse=True)[:5]
    for m in top_measures:
        doc.add_paragraph(
            f"{m['id']}. {m['measure']} (~{m['saving_pct']}%, ≈{m['estimated_annual_rub']:,.0f} ₽/год) — {m['category']}",
            style="List Bullet",
        )
    doc.add_paragraph(
        "Методология: baseline OPEX → целевой WC → NPV/payback → мероприятия по категориям → KPI мониторинг."
    )

    doc.add_heading("Оценка синтеза (хемоинформатика)", level=2)
    from vodopritok.synthesis_assessment import build_synthesis_assessment
    synth = build_synthesis_assessment(ProjectContext(reservoir=reservoir))
    doc.add_paragraph(synth["reduction_story"])
    for a in synth["assessments"][:3]:
        doc.add_paragraph(
            f"#{a['rank']} {a['recipe_id']}: feasibility {a['feasibility_score']}/100 — {a['verdict']}",
            style="List Bullet",
        )

    doc.add_heading("Следующие шаги", level=2)
    for step in [
        "Синтез top-5 → core flood @ T_reservoir (lab gate Frrw≥5, Frro≤2)",
        "DoE по lead formulation → выбор 1–2 рецептур",
        "ОПР на скважине-кандидате → мониторинг WC/Qn",
        "Полный пакет 6 deliverables по ТЗ проекта",
    ]:
        doc.add_paragraph(step, style="List Bullet")

    doc.save(out)
    return out
