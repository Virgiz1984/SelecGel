from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from .cheminformatics import RecipeCandidate
from .decision_tree import TechnologyRecommendation
from .economics import OpexAnalysis
from .lab_program import LabResult
from .models import ProjectContext
from .opr_program import OprResult
from .scouting import build_scouting_summary, load_scouting_excerpts, mechanism_technology_matrix, technology_table_rows


def _set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()


def _title_page(doc: Document, title: str, ctx: ProjectContext) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()
    info = [
        f"Заказчик: {ctx.company_name}",
        f"Исполнитель: {ctx.expert_name}",
        f"Месторождение: {ctx.reservoir.field_name}",
        f"Период проекта: {ctx.project_start} — {ctx.project_end}",
        f"Дата документа: {datetime.now().strftime('%d.%m.%Y')}",
    ]
    for line in info:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def generate_scouting_report(ctx: ProjectContext, path: Path | None = None) -> Path:
    summary = build_scouting_summary(ctx)
    out = path or ctx.ensure_output_dir() / "01-analyticheskiy-otchet-ovp.docx"

    doc = Document()
    _set_doc_defaults(doc)
    _title_page(doc, summary["title"], ctx)

    _heading(doc, "Резюме")
    _para(doc, summary["key_conclusion"])
    _para(doc, summary["cheminformatics_summary"])

    _heading(doc, "1. Постановка задачи")
    r = ctx.reservoir
    _para(
        doc,
        f"Цель проекта — разработка эффективной селективной технологии ограничения водопритока "
        f"для месторождения {r.field_name}. Температура пласта {r.temperature_c}°C, "
        f"минерализация {r.salinity_g_l} г/л, проницаемость {r.permeability_md} мД, "
        f"обводнённость {r.water_cut_pct}%.",
    )

    _heading(doc, "2. Обзор технологий")
    _para(doc, f"В базе проанализировано {summary['technologies_count']} технологий, "
               f"из них {summary['selective_count']} селективных.")
    _table(
        doc,
        ["Технология", "Селективность", "Механизм", "T, °C", "Salinity", "TRL", "Класс"],
        technology_table_rows(),
    )

    _heading(doc, "3. Матрица «механизм → технология»")
    _table(
        doc,
        ["Механизм", "1-й выбор", "2-й выбор", "Не рекомендуется"],
        mechanism_technology_matrix(),
    )

    _heading(doc, "4. Хемоинформатика и оценка синтеза рецептур")
    from .synthesis_assessment import build_synthesis_assessment
    synth = build_synthesis_assessment(ctx)
    _para(doc, synth["expert_competency"])
    _para(doc, synth["reduction_story"])
    _table(doc, ["Этап", "Инструмент", "Вход", "Выход"],
           [[s.get("name", ""), s.get("tool", ""), str(s.get("input_count", "")),
             str(s.get("output_count", ""))] for s in synth["pipeline"].get("stages", [])])
    if synth["assessments"]:
        _heading(doc, "4.1. Feasibility top-5 для синтеза", level=2)
        _table(doc, ["Rank", "ID", "Состав", "MW kDa", "Score", "FTO", "Вердикт", "Примечание"],
               [[str(a["rank"]), a["recipe_id"], a["monomers"], str(a["target_mw_kda"]),
                 str(a["feasibility_score"]), a["fto_risk"], a["verdict"], a["notes"][:80]]
                for a in synth["assessments"]])
    _para(doc, synth["conclusion"], bold=True)
    _para(doc, "Методология: " + " → ".join(m["title"] for m in synth["methodology"]))

    _heading(doc, "5. Рекомендации")
    recs: list[TechnologyRecommendation] = summary["recommendations"]
    for rec in recs:
        _para(doc, f"{rec.rank}. {rec.name_ru} ({rec.track}, score={rec.score:.0f})", bold=True)
        _para(doc, rec.rationale)
        for w in rec.warnings:
            _para(doc, f"  • {w}")

    _heading(doc, "6. Структура отчёта (полная версия)")
    for sec in summary["sections"]:
        _para(doc, sec)

    excerpts = load_scouting_excerpts()
    if excerpts:
        _heading(doc, "7. Материалы scouting (фрагменты)")
        for ex in excerpts:
            _heading(doc, ex["title"], level=2)
            _para(doc, ex["text"])
            _para(doc, f"Источник: scouting/{ex['source']}", bold=False)

    if ctx.session_data and ctx.session_data.get("fto_rows"):
        _heading(doc, "8. FTO / патентная оценка top-5")
        _para(doc, "Предварительная оценка freedom-to-operate для кандидатов из screening.")
        fto_rows = ctx.session_data["fto_rows"]
        _table(
            doc,
            ["mol_id", "Patent", "Класс", "Risk", "Рекомендация"],
            [[r["mol_id"], r["patent_ref"], r["class"], r["risk"], r["recommendation"]] for r in fto_rows],
        )

    doc.save(out)
    return out


def generate_synthesis_assessment_doc(ctx: ProjectContext, assessment: dict, path: Path | None = None) -> Path:
    out = path or ctx.ensure_output_dir() / "01b-ocenka-sintez-khemoinformatika.docx"
    doc = Document()
    _set_doc_defaults(doc)
    _title_page(doc, assessment["title"], ctx)

    _heading(doc, "1. Компетенция и цель")
    _para(doc, assessment["expert_competency"])
    _para(doc, assessment["reduction_story"])
    _para(doc, assessment["conclusion"])

    _heading(doc, "2. Методология хемоинформатики")
    for step in assessment["methodology"]:
        _para(doc, f"Шаг {step['step']}. {step['title']}: {step['detail']}")

    pipe = assessment["pipeline"]
    _heading(doc, "3. Конвейер in silico")
    _table(doc, ["Показатель", "Значение"], [
        ["Вход патентная библиотека", str(pipe["total_input"])],
        ["После QSPR", str(pipe["after_qspr"])],
        ["Top-N для синтеза", str(pipe["top5"])],
        ["Сокращение перебора", f"{pipe['reduction_pct']:.0f}%"],
    ])
    if pipe.get("stages"):
        _table(doc, ["Этап", "Tool", "In", "Out"],
               [[s.get("name", ""), s.get("tool", ""), str(s.get("input_count", "")),
                 str(s.get("output_count", ""))] for s in pipe["stages"]])

    _heading(doc, "4. Оценка возможности синтеза (top-5)")
    _table(doc, ["#", "ID", "Monomers", "MW", "Score", "FTO", "Priority", "Verdict", "Notes"],
           [[str(a["rank"]), a["recipe_id"], a["monomers"], str(a["target_mw_kda"]),
             str(a["feasibility_score"]), a["fto_risk"], a["lab_priority"], a["verdict"], a["notes"]]
            for a in assessment["assessments"]])

    _heading(doc, "5. Доступность мономеров")
    _table(doc, ["ID", "Название", "Заряд", "Поставка", "Fit"],
           [[m["id"], m["name_ru"], m["charge"], m["supply"], m["fit_note"]]
            for m in assessment["monomer_supply"]])

    _heading(doc, "6. DoE и следующий шаг")
    _para(doc, f"DoE Phase 2: {assessment['doe_runs']} прогонов. "
               "Lead → lab gate Frrw≥5, Frro≤2 → ОПР.")

    doc.save(out)
    return out


def generate_lab_program_doc(ctx: ProjectContext, program: dict, path: Path | None = None) -> Path:
    out = path or ctx.ensure_output_dir() / "02-programma-laboratornyh-issledovaniy.docx"
    doc = Document()
    _set_doc_defaults(doc)
    _title_page(doc, program["title"], ctx)

    _heading(doc, "1. Цель и задачи")
    _para(doc, "Разработка и лабораторная верификация 1–2 lead formulations селективного ОВП.")
    if ctx.session_data and ctx.session_data.get("pipeline"):
        p = ctx.session_data["pipeline"]
        _para(doc, f"Источник кандидатов: in silico screening ({p.get('total_input', 500)} → top-5).")
    _para(doc, "Стратегия: параллельно Track 1 (RPM, primary) и Track 2 (термотропный gel, backup).")

    brine = program.get("brine_spec", {})
    if brine:
        _heading(doc, "2. Условия пласта и модельной среды")
        _para(doc, f"Месторождение: {brine.get('field', '—')}, скважина {brine.get('well', '—')}.")
        _para(doc, brine.get("brine_recipe", ""))
        _para(doc, brine.get("core_flood_note", ""))
        _table(doc, ["Параметр", "Значение"], [
            ["T, °C", f"{brine.get('temperature_c', '—')}"],
            ["Salinity, г/л", f"{brine.get('salinity_g_l', '—')}"],
            ["k, мД", f"{brine.get('permeability_md', '—')}"],
            ["Литология", str(brine.get("lithology", "—"))],
            ["API", f"{brine.get('api_gravity', '—')}"],
        ])

    _heading(doc, "3. Треки R&D")
    tracks = program["tracks"]
    for key, rec in tracks.items():
        if rec:
            _para(doc, f"{key}: {rec.name_ru} — {rec.rationale}")

    _heading(doc, "4. Track 1 — кандидаты из screening (Top-5)")
    candidates: list[RecipeCandidate] = program["candidates"]
    rows = []
    for c in candidates:
        ratios = ", ".join(f"{k}:{v:.0%}" for k, v in c.monomer_ratios.items())
        rows.append([
            c.recipe_id, c.name_ru, ratios, str(c.target_mw_kda),
            f"{c.concentration_pct}%", f"{c.predicted_frrw}", f"{c.predicted_frro}", str(c.rank),
        ])
    _table(doc, ["ID", "Название", "Monomers", "MW kDa", "Conc.", "Frrw*", "Frro*", "Rank"], rows)
    _para(doc, "* Predicted values (in silico), subject to core flood validation.")

    track2: list[RecipeCandidate] = program.get("track2_backup", [])
    if track2:
        _heading(doc, "5. Track 2 — backup (параллельный синтез)")
        t2_rows = []
        for c in track2:
            ratios = ", ".join(f"{k}:{v:.0%}" for k, v in c.monomer_ratios.items())
            t2_rows.append([c.recipe_id, c.name_ru, ratios, f"{c.concentration_pct}%", c.hypothesis or "—"])
        _table(doc, ["ID", "Название", "Состав", "Conc.", "Гипотеза"], t2_rows)

    queue = program.get("synthesis_queue", [])
    if queue:
        _heading(doc, "6. Очередь синтеза")
        q_rows = [[q["track"], q["id"], q["name"], q["composition"], q["priority"], q["steps"]] for q in queue]
        _table(doc, ["Track", "ID", "Название", "Состав", "Приоритет", "Этапы"], q_rows)

    _heading(doc, "7. Фазы работ")
    for phase in program["phases"]:
        _heading(doc, phase["name"], level=2)
        for task in phase["tasks"]:
            _para(doc, f"• {task}")
        _para(doc, f"Gate: {phase['gate']}", bold=True)

    gate = program.get("gate_criteria", [])
    if gate:
        _heading(doc, "8. Критерии gate (RPM)")
        _table(doc, ["Метрика", "Phase 1", "Phase 2", "Примечание"],
               [[g["metric"], g["phase1"], g["phase2"], g["note"]] for g in gate])

    _heading(doc, "9. Программа испытаний")
    test_rows = [[t["id"], t["name"], t["method"], str(t["samples"])] for t in program["tests"]]
    _table(doc, ["ID", "Испытание", "Метод", "N samples"], test_rows)

    doe = program.get("doe_runs", [])[:12]
    if doe:
        _heading(doc, "10. DoE-матрица (Phase 2, фрагмент)")
        doe_keys = list(doe[0].keys())
        doe_rows = [[str(run.get(k, "")) for k in doe_keys] for run in doe]
        _table(doc, doe_keys, doe_rows)

    timeline = program.get("timeline", [])
    if timeline:
        _heading(doc, "11. Календарный план (8 недель)")
        _table(doc, ["Нед.", "Track 1", "Track 2", "Milestone"],
               [[t["week"], t["track1"], t["track2"], t["milestone"]] for t in timeline])

    progress = program.get("lab_progress", {})
    if progress.get("has_data"):
        _heading(doc, "12. Статус lab data (QSPRpred / CSV)")
        _para(doc, f"Источник: {progress.get('source', '—')}. Gate пройден: {progress['passed']}/{progress['total']}.")
        if progress.get("rows"):
            _table(doc, ["ID", "Frrw", "Frro", "Oil regain", "Water regain", "Gate"],
                   [[r["recipe_id"], r["frrw"], r["frro"], r["oil_regain"], r["water_regain"], r["gate"]]
                    for r in progress["rows"]])

    _heading(doc, "13. Оборудование")
    for eq in program["equipment"]:
        _para(doc, f"• {eq}")

    _heading(doc, "14. Deliverables программы")
    for item in program.get("deliverables", []):
        _para(doc, f"• {item}")

    doc.save(out)
    return out


def generate_lab_report_doc(ctx: ProjectContext, report: dict, path: Path | None = None) -> Path:
    out = path or ctx.ensure_output_dir() / "03-otchet-laboratoriya.docx"
    doc = Document()
    _set_doc_defaults(doc)
    _title_page(doc, report["title"], ctx)

    _heading(doc, "1. Сводка")
    _para(doc, f"Протестировано рецептур: {report['total_count']}. Gate пройден: {report['passed_count']}.")
    _para(doc, report["conclusion"])

    _heading(doc, "2. Результаты core flood")
    results: list[LabResult] = report["results"]
    rows = []
    for r in results:
        rows.append([
            r.recipe_id, r.recipe_name, f"{r.frrw:.1f}", f"{r.frro:.1f}",
            f"{r.oil_regain_pct:.0f}%", f"{r.water_regain_pct:.0f}%",
            "ДА" if r.passed_gate else "НЕТ",
        ])
    _table(doc, ["ID", "Рецептура", "Frrw", "Frro", "Oil regain", "Water regain", "Gate"], rows)

    _heading(doc, "3. Lead formulations")
    for r in report["lead_recipes"]:
        _para(doc, f"• {r.recipe_name} ({r.recipe_id}): Frrw={r.frrw:.1f}, Frro={r.frro:.1f}")

    if ctx.session_data and ctx.session_data.get("qsprpred"):
        val = ctx.session_data["qsprpred"]
        rep = val.get("report")
        if rep and rep.rows:
            _heading(doc, "4. QSPRpred validation (predicted vs observed)")
            _para(doc, rep.summary)
            q_rows = []
            for row in rep.rows:
                if row.property_name == "frrw":
                    q_rows.append([
                        row.mol_id, row.property_name,
                        f"{row.predicted:.2f}", f"{row.observed:.2f}",
                        f"{row.abs_error:.2f}", f"{row.pct_error:.1f}%",
                    ])
            if q_rows:
                _table(doc, ["mol_id", "Property", "Predicted", "Observed", "|Error|", "%Error"], q_rows)
            if ctx.session_data.get("lab_source"):
                _para(doc, f"Lab data source: {ctx.session_data['lab_source']}")

    doc.save(out)
    return out


def generate_opr_program_doc(ctx: ProjectContext, program: dict, path: Path | None = None) -> Path:
    out = path or ctx.ensure_output_dir() / "04-programma-opr.docx"
    doc = Document()
    _set_doc_defaults(doc)
    _title_page(doc, program["title"], ctx)

    _heading(doc, "1. Цель и компетенция")
    _para(doc, program.get("expert_competency", ""))
    rs = program.get("reservoir_summary", {})
    _para(doc, (
        f"Месторождение {program['field']}, скважина {program['candidate_well']}. "
        f"T={rs.get('temperature_c')}°C, WC={rs.get('water_cut_pct')}%, Qn={rs.get('oil_rate_tpd')} t/d, "
        f"механизм: {rs.get('mechanism')}."
    ))

    lab = program.get("lab_gate", {})
    _heading(doc, "2. Предпосылка: lab gate")
    _para(doc, f"Lead: {lab.get('lead_id', '—')}. Gate: {lab.get('passed', 0)}/{lab.get('total', 0)}. {lab.get('note', '')}")

    _heading(doc, "3. Скважина-кандидат")
    _para(doc, f"Candidate score: {program['well_score']:.0f}/100")
    for note in program["well_notes"]:
        _para(doc, f"• {note}")
    _table(doc, ["Скважина", "Score", "Роль"],
           [[w["name"], f"{w['score']:.0f}", w.get("role", "—")] for w in program["wells_shortlist"]])

    _heading(doc, "4. Технология и lead formulation")
    _para(doc, f"{program['technology']} ({program.get('track', '')})")
    inj = program["injection"]
    _table(doc, ["Параметр", "Значение"], [
        ["Lead mol_id", inj["lead_mol_id"]],
        ["Метод", inj["method"]],
        ["Концентрация", f"{inj['concentration_pct']}%"],
        ["Объём", f"{inj['volume_m3']} m³"],
        ["Скорость", f"{inj['injection_rate_m3h']} m³/ч"],
        ["Длительность", f"{inj['duration_h']} ч"],
        ["Flush", f"{inj['flush_m3']} m³"],
        ["P max", f"{inj['max_pressure_atm']} atm"],
        ["Раствор", inj["brine_match"]],
    ])

    _heading(doc, "5. Фазы ОПР")
    for phase in program["phases"]:
        _heading(doc, phase["phase"], level=2)
        for act in phase["activities"]:
            _para(doc, f"• {act}")

    _heading(doc, "6. Календарный план")
    _table(doc, ["Нед.", "Фаза", "Работы"],
           [[t["week"], t["phase"], t["activities"]] for t in program.get("timeline", [])])

    _heading(doc, "7. Мониторинг")
    _table(doc, ["Период", "Метрики", "Частота"],
           [[m["period"], m["metrics"], m["frequency"]] for m in program.get("monitoring_schedule", [])])

    _heading(doc, "8. KPI и gate ОПР")
    kpi_rows = [[k["metric"], k["target"]] for k in program["kpis"]]
    _table(doc, ["Метрика", "Целевое значение"], kpi_rows)
    gate = program.get("opr_gate", [])
    if gate:
        _table(doc, ["Критерий", "Цель", "При fail"],
               [[g["criterion"], g["target"], g["action_fail"]] for g in gate])

    eco = program.get("economics_preview", {})
    if eco:
        _heading(doc, "9. Экономический прогноз")
        _para(doc, (
            f"WC после: {eco.get('wc_after_pct', '—')}%. "
            f"Экономия: {eco.get('annual_savings_rub', 0):,.0f} ₽/год. "
            f"Payback: {eco.get('payback_months', 0):.1f} мес. "
            f"NPV 3 года: {eco.get('npv_3yr_rub', 0):,.0f} ₽."
        ))

    _heading(doc, "10. HSE и оборудование")
    for item in program.get("hse_checklist", []):
        _para(doc, f"• {item}")
    for eq in program.get("equipment", []):
        _para(doc, f"• {eq}")

    _heading(doc, "11. Contingency и риски")
    for c in program.get("contingency", []):
        _para(doc, f"• {c}")
    for risk in program["risks"]:
        _para(doc, f"• {risk}")

    _heading(doc, "12. Deliverables ОПР")
    for d in program.get("deliverables", []):
        _para(doc, f"• {d}")

    doc.save(out)
    return out


def generate_opr_report_doc(ctx: ProjectContext, report: dict, path: Path | None = None) -> Path:
    out = path or ctx.ensure_output_dir() / "05-otchet-opr.docx"
    doc = Document()
    _set_doc_defaults(doc)
    _title_page(doc, report["title"], ctx)

    _heading(doc, "1. Сводка")
    _para(doc, report["conclusion"])

    _heading(doc, "2. Результаты по скважинам")
    results: list[OprResult] = report["results"]
    rows = []
    for r in results:
        rows.append([
            r.well_name, r.technology, r.treatment_date,
            f"{r.water_cut_before_pct:.0f}%", f"{r.water_cut_after_pct:.0f}%",
            f"{r.oil_rate_before_tpd:.1f}", f"{r.oil_rate_after_tpd:.1f}",
            str(r.effect_duration_months),
        ])
    _table(
        doc,
        ["Скважина", "Технология", "Дата", "WC до", "WC после", "Qn до", "Qn после", "Эффект, мес."],
        rows,
    )

    doc.save(out)
    return out


def generate_opex_plan_doc(ctx: ProjectContext, plan: dict, path: Path | None = None) -> Path:
    out = path or ctx.ensure_output_dir() / "06-plan-snizheniya-opex.docx"
    doc = Document()
    _set_doc_defaults(doc)
    _title_page(doc, plan["title"], ctx)

    analysis: OpexAnalysis = plan["analysis"]
    scenario = plan["scenario"]

    _heading(doc, "1. Контекст и технология")
    _para(doc, f"Месторождение: {ctx.reservoir.field_name}. Технология: {plan.get('technology', 'ОВП')}.")
    _para(doc, f"WC до/после: {scenario.water_cut_before_pct:.0f}% → {scenario.water_cut_after_pct:.0f}%. Qn: {scenario.oil_rate_tpd:.1f} t/d.")
    _para(doc, plan.get("expert_competency", ""))

    methodology = plan.get("methodology", [])
    if methodology:
        _heading(doc, "2. Методология разработки мероприятий по OPEX")
        for step in methodology:
            _para(doc, f"Шаг {step['step']}. {step['title']}: {step['detail']}")

    tech_cmp = plan.get("technology_comparison", [])
    if tech_cmp:
        _heading(doc, "3. Экономика классов технологий ОВП")
        _table(doc, ["Технология", "Track", "CAPEX", "OPEX реагента/год", "Риск fail", "Применимость"],
               [[t["technology"], t["track"], f"{t['capex_rub']:,.0f} ₽", f"{t['annual_reagent_rub']:,.0f} ₽",
                 f"{t['failed_risk_pct']}%", t["best_for"]] for t in tech_cmp])

    _heading(doc, "4. Экономический эффект от сокращения водопритока")
    _table(
        doc,
        ["Показатель", "Значение"],
        [
            ["OPEX воды до, ₽/год", f"{analysis.baseline_water_opex_rub:,.0f}"],
            ["OPEX воды после, ₽/год", f"{analysis.water_opex_after_rub:,.0f}"],
            ["Сокращение воды, m³/год", f"{analysis.water_reduction_m3_year:,.0f}"],
            ["Прямая годовая экономия, ₽", f"{analysis.annual_savings_rub:,.0f}"],
            ["Стоимость обработки, ₽", f"{analysis.treatment_cost_rub:,.0f}"],
            ["Чистый годовой эффект, ₽", f"{analysis.net_annual_benefit_rub:,.0f}"],
            ["Payback, мес.", f"{analysis.payback_months:.1f}"],
            ["₽/ton сокращённой воды", f"{analysis.cost_per_ton_reduced_water_rub:,.0f}"],
            ["NPV 3 года, ₽", f"{analysis.npv_3yr_rub:,.0f}"],
        ],
    )

    breakdown = plan.get("cost_breakdown", [])
    if breakdown:
        _heading(doc, "5. Структура OPEX воды")
        _table(doc, ["Статья", "До, ₽/год", "После, ₽/год"],
               [[b["item"], f"{b['before_rub']:,.0f}", f"{b['after_rub']:,.0f}"] for b in breakdown])

    _heading(doc, "6. Мероприятия по снижению OPEX")
    measure_rows = [
        [
            str(m["id"]), m["category"], m["measure"],
            f"~{m['saving_pct']}%", f"{m.get('estimated_annual_rub', 0):,.0f} ₽",
            m["horizon"], m["owner"], m["description"],
        ]
        for m in plan.get("measures", analysis.measures)
    ]
    _table(doc, ["№", "Категория", "Мероприятие", "Экономия", "₽/год*", "Срок", "Ответственный", "Описание"], measure_rows)
    _para(doc, "* Оценочный дополнительный эффект поверх прямой экономии воды; не суммируется линейно без верификации.")

    _heading(doc, "7. Сводный потенциал")
    _para(doc, f"Прямая экономия воды: {analysis.annual_savings_rub:,.0f} ₽/год.")
    _para(doc, f"Доп. мероприятия (оценка): {plan.get('incremental_savings_rub', 0):,.0f} ₽/год.")
    _para(doc, f"Итого потенциал (ориентир): {plan.get('total_potential_rub', analysis.annual_savings_rub):,.0f} ₽/год.")
    _para(doc, plan.get("portfolio_note", ""))

    kpi = plan.get("monitoring_kpi", [])
    if kpi:
        _heading(doc, "8. KPI мониторинга после внедрения")
        _table(doc, ["KPI", "Цель", "Периодичность"],
               [[k["kpi"], k["target"], k["frequency"]] for k in kpi])

    lifecycle = plan.get("implementation_lifecycle", [])
    if lifecycle:
        _heading(doc, "9. Экономика жизненного цикла внедрения")
        _table(doc, ["Фаза", "Затраты, ₽", "Эффект, ₽/год"],
               [[l["phase"], f"{l['cost_rub']:,.0f}", f"{l['benefit_rub']:,.0f}"] for l in lifecycle])

    _heading(doc, "10. Roadmap внедрения")
    for item in plan["implementation_roadmap"]:
        _para(doc, f"Месяц {item['month']}: {item['action']}")

    doc.save(out)
    return out


def generate_all_deliverables(ctx: ProjectContext | None = None) -> list[Path]:
    from .demo.context_builder import build_deliverable_context
    from .economics import build_opex_plan
    from .lab_program import build_lab_program, build_lab_report
    from .opr_program import build_opr_program, build_opr_report

    ctx = ctx or build_deliverable_context()
    lab_prog = build_lab_program(ctx)
    lab_rep = build_lab_report(ctx)
    opr_prog = build_opr_program(ctx)
    opr_rep = build_opr_report(ctx)
    opex = build_opex_plan(ctx)

    return [
        generate_scouting_report(ctx),
        generate_lab_program_doc(ctx, lab_prog),
        generate_lab_report_doc(ctx, lab_rep),
        generate_opr_program_doc(ctx, opr_prog),
        generate_opr_report_doc(ctx, opr_rep),
        generate_opex_plan_doc(ctx, opex),
    ]
